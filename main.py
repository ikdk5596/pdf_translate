# Extract pdf infomation for translate pdfs.
# Author : sanghong kim
# Organization : Repeblic of Korean Army
# Reference : https://realpython.com/pdf-python/
# 이후 Regular expression을 사용하여 replace 대체

import re
import PyPDF2
# from pdf2docx import parse
from CustomPdf2Docx.CustomPdf2Docx import parse
# from pdf2docx 
import docx
import time
import argparse
import layout_scanner
from pdfminer.pdfpage import PDFPage
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
import operator
from functools import reduce
import sys
from google_trans_new import google_translator


def mergeListAlternative(lst1: list, lst2: list) -> list:
    return list(reduce(operator.add, zip(lst1, lst2)))


def getText(filename: str) -> list:
    doc = Document(filename)
    return [para.text.strip() for para in doc.paragraphs]
  

def progressBar(iterable, task):
    total = len(iterable)
    def printProgressBar(iteration):
        percent = ("{0:.1f}").format(100 * (iteration / float(total)))
        filledLength = int(50 * iteration // total)
        bar = '█' * filledLength + '-' * (50 - filledLength)
        print(f'\rProgress {task}: |{bar}| {percent}% Complete', end='\r')
    printProgressBar(0)
    for i, item in enumerate(iterable, 1):
        yield item
        printProgressBar(i)
    print()
  
  
def main(args):
  input_path = args.input
  docx_path = 'temp.docx'
  output_path = args.output
  sourceLanguageCode = args.sourceLanguageCode
  targetLanguageCode = args.targetLanguageCode
 
  
  with open(input_path, 'rb') as pdf_file:
    read_pdf = PyPDF2.PdfFileReader(pdf_file)
    number_of_pages = read_pdf.getNumPages()
  
  # convert pdf 2 docx
  parse(input_path, docx_path, start = 0, end = number_of_pages)
  extractedText = getText(docx_path)
  filtered_orig_txt = list(filter(None, extractedText))
  translator = google_translator()
  translated_txt = [translator.translate(txt, lang_src=sourceLanguageCode,
                                         lang_tgt=targetLanguageCode) for txt in progressBar(filtered_orig_txt, 'translate')]
  combined_txt = mergeListAlternative(filtered_orig_txt, translated_txt)
  
  output_doc = Document()
  styles = output_doc.styles
  styleE = styles.add_style('English', WD_STYLE_TYPE.PARAGRAPH)
  styleE.font.name = 'Times New Roman'
  styleE.font.bold = True
  styleE.font.size = Pt(11)
  styleK = styles.add_style('Korean', WD_STYLE_TYPE.PARAGRAPH)
  styleK.font.name = 'Times New Roman'
  styleK.font.bold = False
  styleK.font.size = Pt(11)
  for i, txt in enumerate(progressBar(combined_txt, 'Export')):
      p = output_doc.add_paragraph(txt)
      paragraph_format = p.paragraph_format
      paragraph_format.line_spacing = 1
      p.style = output_doc.styles['English'] if i % 2 else output_doc.styles['Korean']
  output_doc.save(output_path)

  # translation = translator.translate('this is a test', src='en', dest='ko')
  # translation = translator.translate('Hello', lang_tgt='ko')
  # print(translation)


if __name__ == '__main__':
  parser = argparse.ArgumentParser(description = 'PDF to TEXT converter',
                                  prog='python -m pdf_layout_scanner')
  parser.add_argument('--input', type=str, dest='input', default='sample.pdf',
                     help='input PDF file')
  parser.add_argument('--output', type=str, dest='output', default='output.docx',
                      help='output PDF file')
  parser.add_argument('--SLC', type=str, dest='sourceLanguageCode', default='en',
                      help="Source LanguageCode")
  parser.add_argument('--TLC', type=str, dest='targetLanguageCode', default='ko',
                      help="Target LanguageCode")
  args = parser.parse_args()
  
  main(args)
  # path = 'sample.pdf'
  # txt, information = extract_information(path)
  # print(txt)
